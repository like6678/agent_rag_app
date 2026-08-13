"""
文档加载器
支持: PDF / TXT / Markdown / DOCX
"""
import os
from typing import List
from loguru import logger

from pypdf import PdfReader
from docx import Document as DocxDocument
import markdown
from bs4 import BeautifulSoup


class DocumentLoader:
    """多格式文档加载,统一输出纯文本"""

    @staticmethod
    def load(file_path: str) -> str:
        """根据扩展名自动选择加载器"""
        ext = os.path.splitext(file_path)[1].lower()
        loaders = {
            ".pdf": DocumentLoader._load_pdf,
            ".txt": DocumentLoader._load_txt,
            ".md": DocumentLoader._load_markdown,
            ".docx": DocumentLoader._load_docx,
        }
        loader = loaders.get(ext)
        if loader is None:
            raise ValueError(f"不支持的文件格式: {ext}")
        text = loader(file_path)
        logger.info(f"加载文档 {file_path}: {len(text)} 字符")
        return text

    @staticmethod
    def load_bytes(content: bytes, ext: str) -> str:
        """从字节加载(用于 MinIO 下载后的内容)"""
        ext = ext.lower()
        if ext == ".pdf":
            return DocumentLoader._load_pdf_bytes(content)
        if ext == ".txt":
            return content.decode("utf-8", errors="ignore")
        if ext == ".md":
            return DocumentLoader._load_markdown_text(content.decode("utf-8", errors="ignore"))
        if ext == ".docx":
            return DocumentLoader._load_docx_bytes(content)
        raise ValueError(f"不支持的文件格式: {ext}")

    @staticmethod
    def _load_pdf(path: str) -> str:
        reader = PdfReader(path)
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    @staticmethod
    def _load_pdf_bytes(content: bytes) -> str:
        import io
        reader = PdfReader(io.BytesIO(content))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    @staticmethod
    def _load_txt(path: str) -> str:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    @staticmethod
    def _load_markdown(path: str) -> str:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return DocumentLoader._load_markdown_text(f.read())

    @staticmethod
    def _load_markdown_text(text: str) -> str:
        """Markdown -> HTML -> 纯文本"""
        html = markdown.markdown(text)
        soup = BeautifulSoup(html, "html.parser")
        return soup.get_text(separator="\n")

    @staticmethod
    def _load_docx(path: str) -> str:
        doc = DocxDocument(path)
        return "\n".join(p.text for p in doc.paragraphs)

    @staticmethod
    def _load_docx_bytes(content: bytes) -> str:
        import io
        doc = DocxDocument(io.BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs)


document_loader = DocumentLoader()
